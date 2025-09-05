import os
import docx
import fitz  # PyMuPDF
import json
import language_tool_python
from groq import Groq
from io import BytesIO
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Use Groq LLM client as it is open source and free to use
try:
    client = Groq()
except Exception as e:
    print(f"Error initializing Groq client: {e}")

# English writing guidelines for compliance checks and modifications
GUIDELINES = """
1.  **Clarity**: Sentences should be clear, concise, and unambiguous. Avoid jargon and complex sentence structures.
2.  **Tone**: Maintain a professional and formal tone suitable for a business or academic context.
3.  **Structure**: Ensure paragraphs are well-structured, focusing on a single topic. Vary sentence length to improve readability.
4.  **Active Voice**: Prefer active voice over passive voice for more direct and engaging writing.
"""

def extract_text(file_contents: bytes, filename: str) -> str:
    """Extracts text from PDF or DOCX file contents."""
    if filename.endswith(".docx"):
        doc = docx.Document(BytesIO(file_contents))
        return "\n".join([para.text for para in doc.paragraphs])
    elif filename.endswith(".pdf"):
        doc = fitz.open(stream=file_contents, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        return text
    else:
        raise ValueError("Unsupported file type")

def take_llm_feedback(text: str) -> dict:
    """Uses an LLM to check for nuanced issues like clarity and tone."""
    prompt = f"""
    You are an expert English editor. Analyze the following text based on these guidelines:
    ---GUIDELINES---
    {GUIDELINES}
    ---TEXT---
    {text[:4000]} # Analyze first 4000 characters to manage token limits

    Provide a brief analysis of the text's adherence to the english guidelines and give a clarity score from 1-10.
    Respond in JSON format with two keys: "score" (integer) and "analysis" (string).
    """
    try:
        """
        1. used model llama-3.3-70b-versatile, as it has higher context length and better performance
        2. used model gemma2-9b-it, capable of understanding complex rules, tone, and clarity, faster and cost-effective
        """
        response = client.chat.completions.create(
            model=os.getenv("GROQ_MODEL"),
            messages=[{"role": "system", "content": "You are an expert editor."},
                      {"role": "user", "content": prompt}],
            response_format={"type": "json_object"}
        )
        
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        print(f"LLM analysis failed: {e}")
        return {"score": "N/A", "analysis": "LLM analysis could not be performed."}

def check_document_rules(text: str) -> dict:
    """Checks text against guidelines using LanguageTool and an LLM."""
    tool = language_tool_python.LanguageTool('en-US')
    matches = tool.check(text)

    grammar_errors = [
        {
            "rule": match.ruleId,
            "message": match.message,
            "context": match.context,
            "replacements": match.replacements,
        }
        for match in matches
    ]

    clarity_feedback = take_llm_feedback(text)

    return {
        "summary": {
            "issue_count": len(grammar_errors),
            "clarity_score": clarity_feedback.get("score", "N/A")
        },
        "details": grammar_errors,
        "ai_analysis": clarity_feedback.get("analysis", "")
    }

def modify_document_with_agent(text: str) -> str:
    """Uses an LLM to rewrite the text according to the defined guidelines."""
    prompt = f"""
    You are an expert document editor. Rewrite the following text to be fully compliant
    with the english guidelines provided. Fix all grammar, spelling, clarity, and tone issues.
    Preserve the original meaning and intent of the text. Return only the rewritten text, without any additional comments or introductions.
    ---GUIDELINES---
    {GUIDELINES}
    ---ORIGINAL TEXT---
    {text}
    ---MODIFIED TEXT---
    """
    try:
        response = client.chat.completions.create(
            model=os.getenv("GROQ_MODEL"),
            messages=[{"role": "user", "content": prompt}]
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print(f"Modification failed: {e}")
        return "Failed to modify the document."

def create_docx_from_text(text: str) -> BytesIO:
    """Creates a docx file in memory from a string of text."""
    doc = docx.Document()
    doc.add_paragraph(text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer